from __future__ import annotations

import csv
import re
from dataclasses import asdict, dataclass
from io import BytesIO, StringIO
from pathlib import Path
from typing import Any

from charset_normalizer import from_bytes
from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".csv", ".txt", ".md"}
UNSUPPORTED_OFFICE_EXTENSIONS = {".doc", ".xls", ".ppt", ".pptx"}


class DocumentParseError(Exception):
    """Base exception for document parsing failures."""


class UnsupportedDocumentTypeError(DocumentParseError):
    """Raised when a file extension is not supported by the first ingestion version."""


class EmptyDocumentError(DocumentParseError):
    """Raised when a supported document contains no extractable text."""


@dataclass(slots=True)
class ParsedElement:
    type: str
    text: str
    metadata: dict[str, Any]

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)

    @classmethod
    def from_dict(cls, value: dict[str, Any]) -> "ParsedElement":
        return cls(
            type=str(value.get("type") or "text"),
            text=str(value.get("text") or ""),
            metadata=dict(value.get("metadata") or {}),
        )


def parse_document_bytes(content: bytes, filename: str, content_type: str | None = None) -> list[ParsedElement]:
    extension = Path(filename).suffix.lower()
    if extension in UNSUPPORTED_OFFICE_EXTENSIONS:
        raise UnsupportedDocumentTypeError(
            f"暂不支持 {extension} 格式，请转换为 PDF、DOCX、XLSX、CSV、TXT 或 MD 后上传。"
        )
    if extension not in SUPPORTED_EXTENSIONS:
        raise UnsupportedDocumentTypeError(
            "暂不支持该文件格式。当前支持 PDF、DOCX、XLSX、CSV、TXT、MD。"
        )
    if not content:
        raise EmptyDocumentError("上传文件为空。")

    try:
        if extension == ".pdf":
            elements = _parse_pdf(content, filename)
        elif extension == ".docx":
            elements = _parse_docx(content, filename)
        elif extension == ".xlsx":
            elements = _parse_xlsx(content, filename)
        elif extension == ".csv":
            elements = _parse_csv(content, filename)
        else:
            elements = parse_text_content(_decode_text(content), filename, parser=extension.lstrip("."))
    except DocumentParseError:
        raise
    except Exception as exc:  # noqa: BLE001
        raise DocumentParseError(f"文档解析失败：{exc}") from exc

    if not any(element.text.strip() for element in elements):
        if extension == ".pdf":
            raise EmptyDocumentError("PDF 未提取到可复制文本，扫描版 PDF 暂不支持 OCR。")
        raise EmptyDocumentError("文档未提取到可索引文本。")
    return elements


def parse_text_content(raw_text: str, source_name: str | None = None, parser: str = "plain_text") -> list[ParsedElement]:
    text = raw_text.replace("\r\n", "\n").replace("\r", "\n").strip()
    if not text:
        return []

    elements: list[ParsedElement] = []
    section_path: list[str] = []
    blocks = [block.strip() for block in re.split(r"\n\s*\n", text) if block.strip()]
    for block in blocks:
        heading_match = re.match(r"^(#{1,6})\s+(.+)$", block)
        metadata = {
            "parser": parser,
            "source_name": source_name,
        }
        if heading_match:
            level = len(heading_match.group(1))
            heading = heading_match.group(2).strip()
            section_path = section_path[: level - 1] + [heading]
            metadata["section_path"] = section_path.copy()
            elements.append(ParsedElement(type="heading", text=heading, metadata=metadata))
            continue
        if section_path:
            metadata["section_path"] = section_path.copy()
        elements.append(ParsedElement(type="text", text=block, metadata=metadata))
    return elements


def elements_to_text(elements: list[ParsedElement]) -> str:
    return "\n\n".join(element.text.strip() for element in elements if element.text.strip())


def _parse_pdf(content: bytes, filename: str) -> list[ParsedElement]:
    reader = PdfReader(BytesIO(content))
    elements: list[ParsedElement] = []
    for page_index, page in enumerate(reader.pages, start=1):
        page_text = (page.extract_text() or "").strip()
        if not page_text:
            continue
        elements.append(
            ParsedElement(
                type="text",
                text=page_text,
                metadata={
                    "parser": "pdf",
                    "source_name": filename,
                    "page_start": page_index,
                    "page_end": page_index,
                },
            )
        )
    return elements


def _parse_docx(content: bytes, filename: str) -> list[ParsedElement]:
    document = Document(BytesIO(content))
    elements: list[ParsedElement] = []
    section_path: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = paragraph.style.name if paragraph.style is not None else ""
        metadata = {
            "parser": "docx",
            "source_name": filename,
        }
        heading_match = re.match(r"Heading\s+(\d+)", style_name or "")
        if heading_match:
            level = max(int(heading_match.group(1)), 1)
            section_path = section_path[: level - 1] + [text]
            metadata["section_path"] = section_path.copy()
            elements.append(ParsedElement(type="heading", text=text, metadata=metadata))
        else:
            if section_path:
                metadata["section_path"] = section_path.copy()
            elements.append(ParsedElement(type="text", text=text, metadata=metadata))

    for table_index, table in enumerate(document.tables, start=1):
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        rows = [row for row in rows if any(cell for cell in row)]
        if not rows:
            continue
        headers = rows[0]
        data_rows = [{"index": idx, "values": row} for idx, row in enumerate(rows[1:], start=2)]
        metadata = {
            "parser": "docx",
            "source_name": filename,
            "section_path": section_path.copy(),
            "table_index": table_index,
            "headers": headers,
            "rows": data_rows,
        }
        elements.append(
            ParsedElement(
                type="table",
                text=_table_rows_to_text(headers, data_rows),
                metadata=metadata,
            )
        )
    return elements


def _parse_xlsx(content: bytes, filename: str) -> list[ParsedElement]:
    workbook = load_workbook(BytesIO(content), read_only=True, data_only=True)
    elements: list[ParsedElement] = []
    for worksheet in workbook.worksheets:
        raw_rows: list[tuple[int, list[str]]] = []
        for row_index, row in enumerate(worksheet.iter_rows(values_only=True), start=1):
            values = [_cell_to_text(value) for value in row]
            if any(values):
                raw_rows.append((row_index, values))
        if not raw_rows:
            continue
        headers = raw_rows[0][1]
        data_rows = [{"index": row_index, "values": values} for row_index, values in raw_rows[1:]]
        metadata = {
            "parser": "xlsx",
            "source_name": filename,
            "sheet_name": worksheet.title,
            "headers": headers,
            "rows": data_rows,
        }
        elements.append(
            ParsedElement(
                type="table",
                text=_table_rows_to_text(headers, data_rows, worksheet.title),
                metadata=metadata,
            )
        )
    workbook.close()
    return elements


def _parse_csv(content: bytes, filename: str) -> list[ParsedElement]:
    text = _decode_text(content)
    reader = csv.reader(StringIO(text))
    raw_rows = [(index, [cell.strip() for cell in row]) for index, row in enumerate(reader, start=1)]
    raw_rows = [(index, row) for index, row in raw_rows if any(row)]
    if not raw_rows:
        return []
    headers = raw_rows[0][1]
    data_rows = [{"index": row_index, "values": values} for row_index, values in raw_rows[1:]]
    sheet_name = Path(filename).stem
    metadata = {
        "parser": "csv",
        "source_name": filename,
        "sheet_name": sheet_name,
        "headers": headers,
        "rows": data_rows,
    }
    return [
        ParsedElement(
            type="table",
            text=_table_rows_to_text(headers, data_rows, sheet_name),
            metadata=metadata,
        )
    ]


def _decode_text(content: bytes) -> str:
    try:
        return content.decode("utf-8")
    except UnicodeDecodeError:
        match = from_bytes(content).best()
        if match is not None:
            return str(match)
        return content.decode("utf-8", errors="replace")


def _cell_to_text(value: Any) -> str:
    if value is None:
        return ""
    return str(value).strip()


def _table_rows_to_text(
    headers: list[str],
    rows: list[dict[str, Any]],
    label: str | None = None,
) -> str:
    lines = []
    if label:
        lines.append(f"表格：{label}")
    if headers:
        lines.append("列：" + " | ".join(headers))
    for row in rows:
        values = [str(value) for value in row.get("values", [])]
        lines.append(f"行 {row.get('index')}: " + _format_row(headers, values))
    return "\n".join(lines)


def _format_row(headers: list[str], values: list[str]) -> str:
    cells = []
    for index, value in enumerate(values):
        header = headers[index] if index < len(headers) and headers[index] else f"列{index + 1}"
        cells.append(f"{header}: {value}")
    return "; ".join(cells)
