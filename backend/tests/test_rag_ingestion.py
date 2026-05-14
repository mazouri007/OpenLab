from __future__ import annotations

from io import BytesIO

import pytest
from docx import Document
from fastapi.testclient import TestClient
from openpyxl import Workbook
from pypdf import PdfWriter
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker

from app.db.base import Base
from app.main import app
from app.models import KnowledgeChunk
from app.schemas.kb import KnowledgeDocumentCreate
from app.services.rag.chunking import chunk_elements
from app.services.rag.bm25 import bm25_scores
from app.services.rag.document_parser import (
    EmptyDocumentError,
    UnsupportedDocumentTypeError,
    parse_document_bytes,
    parse_text_content,
)
from app.services.rag.service import RagService


def test_text_chunker_respects_long_text_limit() -> None:
    elements = parse_text_content("# 规范\n\n" + "这是一段很长的内容。" * 180, "manual")

    chunks = chunk_elements(elements, "研发规范", "manual")

    assert len(chunks) > 1
    assert all(len(chunk.content) <= 1400 for chunk in chunks)
    assert chunks[0].metadata["section_path"] == ["规范"]


def test_bm25_scores_rank_relevant_chinese_chunk() -> None:
    scores = bm25_scores(
        [
            ("api", "治疗方式管理接口支持新增治疗方式、删除治疗方式和年龄段绑定。"),
            ("label", "CVAT 数据标注流程包括创建项目、上传视频和导出标注。"),
            ("generic", "系统支持用户登录和基础页面展示。"),
        ],
        ["新增治疗方式接口"],
    )

    assert scores["api"] == 1.0
    assert scores["api"] > scores["generic"]
    assert scores["api"] > scores["label"]


def test_bm25_scores_return_zero_for_no_keyword_match() -> None:
    scores = bm25_scores([("doc-1", "alpha beta"), ("doc-2", "gamma delta")], ["完全无关"])

    assert scores == {"doc-1": 0.0, "doc-2": 0.0}


def test_text_chunker_keeps_small_heading_with_body() -> None:
    elements = parse_text_content("# 规范\n\n短段落。", "manual")

    chunks = chunk_elements(elements, "研发规范", "manual")

    assert len(chunks) == 1
    assert chunks[0].metadata["chunk_type"] == "text"
    assert "规范" in chunks[0].content
    assert "短段落" in chunks[0].content


def test_csv_parser_and_table_chunker_keep_row_range() -> None:
    content = "name,score\nAlice,95\nBob,88\n".encode("utf-8")
    elements = parse_document_bytes(content, "scores.csv")

    chunks = chunk_elements(elements, "成绩", "scores.csv")

    assert len(elements) == 1
    assert elements[0].type == "table"
    assert chunks[0].metadata["sheet_name"] == "scores"
    assert chunks[0].metadata["row_start"] == 2
    assert chunks[0].metadata["row_end"] == 3
    assert "列：name | score" in chunks[0].content
    assert "name：Alice" in chunks[0].content
    assert chunks[0].content.endswith("。")


def test_table_chunker_avoids_too_few_rows_before_splitting() -> None:
    rows = "\n".join(f"{index},{'长内容' * 160}" for index in range(1, 7))
    elements = parse_document_bytes(f"id,desc\n{rows}\n".encode("utf-8"), "wide.csv")

    chunks = chunk_elements(elements, "宽表", "wide.csv")

    assert len(chunks) > 1
    assert chunks[0].metadata["row_start"] == 2
    assert chunks[0].metadata["row_end"] == 4
    assert chunks[0].metadata["chunk_type"] == "table"


def test_docx_parser_extracts_heading_paragraph_and_table() -> None:
    document = Document()
    document.add_heading("接口规范", level=1)
    document.add_paragraph("所有接口都必须返回统一响应结构。")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "字段"
    table.cell(0, 1).text = "说明"
    table.cell(1, 0).text = "code"
    table.cell(1, 1).text = "业务状态码"
    stream = BytesIO()
    document.save(stream)

    elements = parse_document_bytes(stream.getvalue(), "spec.docx")

    assert [element.type for element in elements] == ["heading", "text", "table"]
    assert elements[1].metadata["section_path"] == ["接口规范"]
    assert elements[2].metadata["headers"] == ["字段", "说明"]


def test_xlsx_parser_creates_structured_table_chunks() -> None:
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "缺陷"
    worksheet.append(["id", "severity"])
    worksheet.append(["BUG-1", "high"])
    worksheet.append(["BUG-2", "low"])
    stream = BytesIO()
    workbook.save(stream)

    elements = parse_document_bytes(stream.getvalue(), "bugs.xlsx")
    chunks = chunk_elements(elements, "缺陷表", "bugs.xlsx")

    assert elements[0].metadata["sheet_name"] == "缺陷"
    assert chunks[0].metadata["row_start"] == 2
    assert chunks[0].metadata["row_end"] == 3
    assert "BUG-1" in chunks[0].content


def test_unsupported_and_empty_documents_are_diagnostic() -> None:
    with pytest.raises(UnsupportedDocumentTypeError):
        parse_document_bytes(b"legacy", "legacy.doc")
    with pytest.raises(EmptyDocumentError):
        parse_document_bytes(b"", "empty.txt")


def test_blank_pdf_reports_ocr_not_supported() -> None:
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    stream = BytesIO()
    writer.write(stream)

    with pytest.raises(EmptyDocumentError, match="扫描版 PDF 暂不支持 OCR"):
        parse_document_bytes(stream.getvalue(), "scan.pdf")


def test_file_upload_rejects_unsupported_format() -> None:
    client = TestClient(app)

    response = client.post(
        "/api/v1/projects/demo-platform/kb/documents/upload",
        files={"file": ("legacy.doc", b"legacy", "application/msword")},
    )

    assert response.status_code == 415
    assert "暂不支持" in response.json()["detail"]


def test_index_document_stores_chunk_identity_hash_and_indexed_content() -> None:
    engine = create_engine("sqlite:///:memory:", future=True)
    Base.metadata.create_all(bind=engine)
    session_factory = sessionmaker(bind=engine, autoflush=False, autocommit=False, future=True)

    with session_factory() as db:
        service = RagService(db)
        service.llm_provider.settings.enable_mock_llm = True
        document = service.create_document(
            "project-1",
            KnowledgeDocumentCreate(
                title="RAG 模块设计",
                source_name="manual",
                raw_text="# 混合检索\n\n## BM25\n\nBM25 主要用于关键词精确匹配。",
            ),
        )
        service.index_document(document)

        chunk = db.query(KnowledgeChunk).filter(KnowledgeChunk.document_id == document.id).one()
        chunk_id = chunk.id
        metadata = chunk.metadata_json
        refreshed_document = db.get(type(document), document.id)
        vector_indexed = refreshed_document.metadata_json["vector_indexed"]
        hits = service.vector_store.query("project-1", [1.0, 0.1, 0.2, 0.3])

    assert metadata["chunk_id"] == f"{document.id}_0001"
    assert metadata["doc_id"] == document.id
    assert metadata["chunk_index"] == 0
    assert len(metadata["content_hash"]) == 32
    assert metadata["chunk_type"] == "manual"
    assert "标题路径：混合检索 > BM25" in metadata["indexed_content"]
    assert "embedding" not in metadata
    assert vector_indexed is True
    assert hits and hits[0].chunk_id == chunk_id


def test_index_document_keeps_keyword_chunks_when_vector_store_unavailable() -> None:
    engine = create_engine("sqlite:///:memory:", future=True)
    Base.metadata.create_all(bind=engine)
    session_factory = sessionmaker(bind=engine, autoflush=False, autocommit=False, future=True)

    with session_factory() as db:
        service = RagService(db)
        service.vector_store = None
        service.llm_provider.settings.enable_mock_llm = True
        document = service.create_document(
            "project-1",
            KnowledgeDocumentCreate(
                title="Fallback",
                source_name="manual",
                raw_text="# 规范\n\n关键词检索仍应可用。",
            ),
        )
        service.index_document(document)

        chunks = db.query(KnowledgeChunk).filter(KnowledgeChunk.document_id == document.id).all()
        refreshed_document = db.get(type(document), document.id)

    assert chunks
    assert refreshed_document.parse_status == "indexed"
    assert refreshed_document.metadata_json["vector_indexed"] is False
    assert "BM25 索引" in refreshed_document.error_message
